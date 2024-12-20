from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor
import settings
import re

"""
В этом файле находятся функции для создания ворд файла. 
Я вынес их в отдельный файл чтобы не загромождать код в файлах страниц приложения.
"""

# Функция форматирования текста задания, которая убирает лишние пробелы и переносы строк.
def text_formater(text):
    text = re.sub(r'\s+', ' ', text)
    return text

# Основная функция создания ворд файла.
def create_docx():
    document = Document()

    # Инициализация.------------------------------------------------------

    first_and_last_name = f"{settings.first_name} {settings.last_name}"

    # ---------------------------------------------------------------------
    # Автор.---------------------------------------------------------------

    # Кто создал документ.
    core_properties = document.core_properties
    core_properties.author = first_and_last_name

    # Кто редактировал документ.
    last_modified_by = document.core_properties
    last_modified_by.last_modified_by = first_and_last_name

    # Комментарий.
    comments = document.core_properties
    comments.comments = " "

    # ---------------------------------------------------------------------
    # Разметка страницы.---------------------------------------------------

    sections = document.sections

    # Цикл для изменения формата каждой страницы.
    for section in sections:
        # Изменения формата листа на A4 для каждой страницы.
        section.page_height = Cm(settings.list_format_height)
        section.page_width = Cm(settings.list_format_width)
        # Изменение полей листа.
        section.top_margin = Cm(settings.list_top_margin)
        section.bottom_margin = Cm(settings.list_bottom_margin)
        section.left_margin = Cm(settings.list_left_margin)
        section.right_margin = Cm(settings.list_right_margin)

    # ---------------------------------------------------------------------
    # Заглавие.------------------------------------------------------------

    main_heading = document.add_heading()
    # Добавление заголовка с типом работы и ее номером.
    run = main_heading.add_run(f"{settings.type_of_work} {settings.work_number}. {text_formater(settings.work_theme)}")

    # Указание стиля шрифта.

    font = run.font
    # Жирный шрифт или нет.
    font.bold = settings.heading_font_bold
    # Название шрифта.
    font.name = settings.heading_font_name
    # Размер шрифта.
    font.size = Pt(settings.heading_font_size)
    # Цвет шрифта.
    font.color.rgb = RGBColor(0, 0, 0)

    # ---------------------------------------------------------------------
    # Создание заданий.----------------------------------------------------

    # Счетчик итераций цикла.
    inter = 0

    # Цикл для создания и форматирования заданий.
    for i in range(settings.start_task - 1, settings.count_of_task):
        # Установка уровня заголовка для задания.
        task_heading = document.add_heading(level=2)
        # Указание номера задания в нужном формате в зависимости от значения.
        if i < 9:
            run = task_heading.add_run("Задание 0" + str(i + 1))
        else:
            run = task_heading.add_run("Задание " + str(i + 1))

        # Указание стиля шрифта.-----------------------------------------------

        font = run.font
        # Жирный шрифт или нет.
        font.bold = settings.task_font_bold
        # Название шрифта.
        font.name = settings.task_font_name
        # Размер шрифта.
        font.size = Pt(settings.task_font_size)
        # Цвет шрифта.
        font.color.rgb = RGBColor(0, 0, 0)

        paragraph_format = task_heading.paragraph_format
        # Добавление отступа для заголовка задания.
        paragraph_format.left_indent = Cm(settings.task_left_indent)

        # ---------------------------------------------------------------------
        # Условие.-------------------------------------------------------------

        condition_paragraph = document.add_paragraph()
        # Разметка текста по ширине.
        condition_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Добавление текста с условием задания.
        run = condition_paragraph.add_run(f"Условие: {text_formater(settings.tasks_text[inter])}")

        # Стиль шрифта.

        font = run.font
        # Название шрифта.
        font.name = settings.condition_font_name
        # Размер шрифта.
        font.size = Pt(settings.condition_font_size)

        # ---------------------------------------------------------------------
        # Описание картинки.---------------------------------------------------

        if settings.checkbox[inter] == False:
            picture_description = document.add_paragraph()
            # Присвоение стиля для текста.
            picture_description.style = settings.picture_description_style
            # Разметка текста по центру.
            picture_description.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Добавление текста подписи рисунка.
            run = picture_description.add_run("Рис " + str(i + 1) + ". ")

        # Стиль шрифта.

        font = run.font
        # Название шрифта.
        font.name = settings.picture_description_name
        # Размер шрифта.
        font.size = Pt(settings.picture_description_size)

        # ---------------------------------------------------------------------

        # Увеличение счетчика итераций цикла на 1.
        inter += 1

    # ---------------------------------------------------------------------
    # Сохранения документа.------------------------------------------------

    # Сохраняет туда же, где находится и сам этот файл
    document.save(f"{settings.type_of_work} {settings.work_number}.docx")
