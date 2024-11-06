# Скачайте python-docx
# powershell --> "py -m pip install python-docx" (без кавычек конечно же вводите эту штуку. Или тупо скопируйте)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor
document = Document()

# Введения
au = input('Введите имя и фамилию автора (т.е. себя): ')
task_type = input("Напишите на русской раскладке з (занятие), п (практическая) или р (работа в классе): ")
number_of_work = input("Введите номер работы в классе/занятия/практической: ")
task_number = int(input("Введите количество заданий: "))
first_task_is = int(input("Введите номер начального задания (например, '00' или '01'): "))

# Автор
core_properties = document.core_properties
core_properties.author = au
last_modified_by = document.core_properties
last_modified_by.last_modified_by = au
comments = document.core_properties
comments.comments= " "

work_type = {"з": "Занятие",
             "п": "Практическая работа",
             "р": "Работа в классе"}
task_type = work_type[task_type]


# Выбор между занятем, практической, работой в классе. Думаю и так всё понятно.
    
# Заглавие
main_heading = document.add_heading()
run = main_heading.add_run((task_type) + " " + (number_of_work))
font = run.font
font.bold = False
font.name = "Arial"
font.size = Pt(20)
font.color.rgb = RGBColor(0, 0, 0)

# Для task_number заданий
for i in range(first_task_is - 1, task_number):
    # Задание
    task_heading = document.add_heading(level = 2)
    if (i < 9):
        run = task_heading.add_run("Задание 0" + str(i + 1))
    else:
        run = task_heading.add_run("Задание " + str(i + 1))
    font = run.font
    font.bold = False
    font.name = "Arial"
    font.size = Pt(18)
    font.color.rgb = RGBColor(0, 0, 0)
    paragraph_format = task_heading.paragraph_format
    paragraph_format.left_indent = Cm(1.5)

    # Условие
    if_paragraph = document.add_paragraph()
    if_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = if_paragraph.add_run("Условие: ")
    font = run.font
    font.name = "Times New Roman"
    font.size = Pt(12)
   
    # Описание картинки. Если не нужно, просто уберёте в ворде
    picture_description = document.add_paragraph()
    picture_description.style = "Quote"
    picture_description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = picture_description.add_run("Рис " + str(i + 1) + ". ")
    font = run.font
    font.name = "Times New Roman"
    font.size = Pt(12)

# Сохраняет туда же, где находится и сам этот файл
document.save((task_type) + " " + (number_of_work) + ".docx")