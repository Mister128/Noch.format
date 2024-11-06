import flet as ft
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor


def main(page: ft.Page):

    # Настройка окна приложения

    page.title="Noch.ka 2.0"
    page.window_always_on_top
    page.window_width = 700
    page.window_height = 630
    page.window_visible = True
    page.window_resizable = False
    page.update()


    # Инициализация переменных

    # Акцентный цвет
    accent_color = "cyan"

    # Фильтр текста
    _filter = ft.InputFilter(regex_string=r"^[0-9]*$")
        
    # Изменение темы
    def theme_changed(e):
        page.theme_mode = (
            ft.ThemeMode.DARK
            if page.theme_mode == ft.ThemeMode.LIGHT
            else ft.ThemeMode.LIGHT
        )
        e.control.selected = not e.control.selected
        e.control.update()
        page.update()
    page.theme_mode = ft.ThemeMode.DARK
    theme = ft.Container(
        ft.IconButton(icon = ft.icons.DARK_MODE, selected_icon=ft.icons.SUNNY,
                    on_click=theme_changed, selected=False),
        alignment=ft.alignment.bottom_right
        )

    # Функция создания текстового документа
    def create_docx(e):
        document = Document()

        # Инициализация
        fln = f"{first_name.content.value} {last_name.content.value}"
        tow = str(type_work.value)
        now = str(number_of_work.content.value)
        cot = int(count_of_task.content.value)
        st = int(start_task.content.value)

        # Автор
        core_properties = document.core_properties
        core_properties.author = fln
        last_modified_by = document.core_properties
        last_modified_by.last_modified_by = fln
        comments = document.core_properties
        comments.comments = " "

        # Заглавие
        main_heading = document.add_heading()
        run = main_heading.add_run((tow) + " " + (now))
        font = run.font
        font.bold = False
        font.name = "Arial"
        font.size = Pt(20)
        font.color.rgb = RGBColor(0, 0, 0)

        # Для task_number заданий
        for i in range(st - 1, cot):
            # Задание
            task_heading = document.add_heading(level=2)
            if i < 9:
                run = task_heading.add_run("Задание 0" + str(i + 1))
            else:
                run = task_heading.add_run("Задание " + str(i + 1))
            font = run.font
            font.bold = False
            font.name = "Arial"
            font.size = Pt(18)
            font.color.rgb = RGBColor(0, 0, 0)
            paragraph_format = task_heading.paragraph_format
            paragraph_format.left_indent = Cm(1.5)

            # Условие
            if_paragraph = document.add_paragraph()
            if_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = if_paragraph.add_run("Условие: ")
            font = run.font
            font.name = "Times New Roman"
            font.size = Pt(12)

            # Описание картинки. Если не нужно, просто уберёте в ворде
            picture_description = document.add_paragraph()
            picture_description.style = "Quote"
            picture_description.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = picture_description.add_run("Рис " + str(i + 1) + ". ")
            font = run.font
            font.name = "Times New Roman"
            font.size = Pt(12)

        # Сохраняет туда же, где находится и сам этот файл
        document.save(tow + " " + now + ".docx")

    # Название
    name = ft.Container(
            ft.Text("Nochka.format", text_align="center", size=50),
            alignment=ft.alignment.top_center,
            bgcolor=accent_color
        )
    
    # Имя и фамилия
    first_name = ft.Container(
            ft.TextField(label="Введите своё имя",
                         border_color=accent_color),
        )

    last_name = ft.Container(
        ft.TextField(label="Введите свою фамилию",
                     border_color=accent_color),
    )
    
    # Текст
    container_variant = ft.Container(
            ft.Text("Выберите вариант работы",
                    size=20),
            margin=ft.margin.only(bottom=-5),
            alignment=ft.alignment.top_center
        )

    # Варианты
    type_work = ft.Dropdown(
        width=700,
        border_color=accent_color,
        hint_text="Выберите тип работы",
        options=[
            ft.dropdown.Option("Занятие"),
            ft.dropdown.Option("Практическая работа"),
            ft.dropdown.Option("Работа в классе"),
        ],
    )

    # Номер работы
    number_of_work = ft.Container(
        ft.TextField(label="Введите номер работы",
                     input_filter=_filter,
                     border_color=accent_color),
        margin=ft.margin.only(top=10),
    )

    # Кол-во заданий
    count_of_task = ft.Container(
        ft.TextField(label="Введите кол-во заданий",
                     input_filter=_filter,
                     border_color=accent_color)
        )

    # Номер начального задания
    start_task = ft.Container(
        ft.TextField(label="Введите номер начального задания",
                     input_filter=_filter,
                     max_length=2,
                     border_color=accent_color),
    )

    # Кнопка создания файла
    create = ft.ElevatedButton("Создать",
                               style=ft.ButtonStyle(shape=ft.StadiumBorder()),
                               color=accent_color,
                               on_click=create_docx,
                               disabled=False)

    # Добавление созданных элементов на страницу
    page.add(name, first_name, last_name, container_variant, type_work, number_of_work, count_of_task, start_task, create, theme)


ft.app(main)